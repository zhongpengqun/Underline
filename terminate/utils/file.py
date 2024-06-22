import sys
from PyPDF2 import PdfReader
from docx import Document

from .clipboard import read_from_clipboard
from .ui import green


def append_clipboard_content_to_specified_text_file(text_file_path):
    clipboard_content = read_from_clipboard()
    if clipboard_content:
        with open(text_file_path, 'a', encoding='utf-8') as f:
            f.write('\n ---------------- \n')
            f.write(clipboard_content)
            print(green(f'Write to {text_file_path} Done!'))

    # # 验证是否写入成功
    # with open(text_file_path, 'rb') as f:
    #     lines = f.readlines()
    #     last_lines = lines[-2:]
    #     print("-------验证是否写入成功---------")
    #     print('\n'.join([x.decode('utf8') for x in last_lines]))
    #     print("-------------------------------")


def remove_debug_sentences():
    pass


def transfer_docx_as_lines(file_path):
    doc = Document(file_path)
    result = []

    for p in doc.paragraphs:
        result.append(p.text)

    for t in doc.tables:
        for column_index in range(len(t.columns)):
            for row_index in range(len(t.rows)):
                cell = t.cell(row_index, column_index)
                result.append(cell.text)
    return result


def transfer_pdf_as_lines(file_path):
    result = []

    reader = PdfReader(file_path)
    for page in reader.pages:
        result.append(page.extract_text())

    return result

